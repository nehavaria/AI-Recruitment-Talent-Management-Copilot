"""
Milestone 3 — Interview Questions Generator
Uses existing DB data (parsed by existing parsers) + Groq API to generate questions.
"""

import html
import json
import streamlit as st
from groq import Groq

from config.settings import GROQ_API_KEY
from services.candidate_service import CandidateService
from ui.components import page_header, empty_state


# ── Groq client (lazy, always reads fresh key) ────────────────────────────

def _groq_client():
    from config.settings import GROQ_API_KEY as _KEY
    return Groq(api_key=_KEY)


# ── Groq generation ────────────────────────────────────────────────────────

def _generate_via_groq(
    job_title: str,
    jd_skills: list[str],
    responsibilities: list[str],
    cand_name: str,
    cand_skills: list[str],
    cand_exp: list[str],
    matched: list[str],
    missing: list[str],
) -> dict:
    """Call Groq and return structured questions dict."""

    prompt = f"""
You are an expert technical interviewer. Generate interview questions for the following:

JOB DETAILS:
- Job Title: {job_title}
- Required Skills: {", ".join(jd_skills) or "Not specified"}
- Responsibilities: {chr(10).join(f"  - {r}" for r in responsibilities[:5]) or "Not specified"}

CANDIDATE DETAILS:
- Name: {cand_name}
- Skills: {", ".join(cand_skills) or "Not specified"}
- Experience: {chr(10).join(f"  - {e}" for e in cand_exp[:4]) or "Not specified"}
- Matched Skills: {", ".join(matched) or "None"}
- Skill Gaps (missing): {", ".join(missing) or "None"}

Generate EXACTLY the following — return ONLY valid JSON, no markdown, no explanation:

{{
  "technical": [
    {{"difficulty": "Easy",   "question": "..."}},
    {{"difficulty": "Easy",   "question": "..."}},
    {{"difficulty": "Medium", "question": "..."}},
    {{"difficulty": "Medium", "question": "..."}},
    {{"difficulty": "Hard",   "question": "..."}}
  ],
  "behavioral": [
    {{"difficulty": "Easy",   "question": "..."}},
    {{"difficulty": "Medium", "question": "..."}},
    {{"difficulty": "Hard",   "question": "..."}}
  ],
  "situational": [
    {{"difficulty": "Medium", "question": "..."}},
    {{"difficulty": "Hard",   "question": "..."}}
  ]
}}

Rules:
- Every question MUST have a "difficulty" field: exactly one of "Easy", "Medium", or "Hard".
- Easy questions test basic knowledge and definitions.
- Medium questions require applied understanding and examples.
- Hard questions require deep expertise, design thinking, or complex problem solving.
- Technical questions must be specific to the job skills and skill gaps.
- Behavioral questions must use the STAR method format.
- Situational questions must present a realistic work scenario.
- Do NOT include candidate name in questions.
- Return ONLY the JSON object above.
"""

    client = _groq_client()
    last_error = None

    for model in ["llama-3.3-70b-versatile", "llama-3.1-8b-instant", "mixtral-8x7b-32768"]:
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
            )
            raw = response.choices[0].message.content.strip()
            if raw.startswith("```"):
                raw = raw.split("```")[1]
                if raw.startswith("json"):
                    raw = raw[4:]
            return json.loads(raw.strip())
        except Exception as e:
            last_error = e
            err = str(e)
            if "429" in err or "rate_limit" in err.lower() or "503" in err:
                continue
            raise

    raise RuntimeError(
        f"All Groq models are unavailable. Please wait a moment and try again.\n{last_error}"
    )


# ── UI helpers ─────────────────────────────────────────────────────────────

def _pills(items: list[str], color: str, bg: str, border: str) -> str:
    return "".join(
        f"<span style='background:{bg};color:{color};border:1px solid {border};"
        f"padding:4px 12px;border-radius:20px;font-size:0.75rem;font-weight:600;"
        f"margin:3px;display:inline-block'>{html.escape(s.strip())}</span>"
        for s in items if s.strip()
    )


def _section_label(text: str, color: str = "#94a3b8") -> None:
    st.markdown(
        f"<div style='font-size:0.68rem;font-weight:700;color:{color};"
        f"text-transform:uppercase;letter-spacing:0.1em;margin-bottom:8px'>"
        f"{text}</div>",
        unsafe_allow_html=True,
    )


def _info_row(label: str, value: str) -> None:
    st.markdown(
        f"<div style='display:flex;gap:12px;padding:8px 0;"
        f"border-bottom:1px solid rgba(255,255,255,0.06)'>"
        f"<div style='font-size:0.75rem;font-weight:700;color:#64748b;min-width:130px'>"
        f"{html.escape(label)}</div>"
        f"<div style='font-size:0.82rem;color:#f1f5f9;font-weight:500'>"
        f"{html.escape(str(value or '—'))}</div></div>",
        unsafe_allow_html=True,
    )


_DIFF_COLORS = {
    "Easy":   ("#10b981", "rgba(16,185,129,0.15)",  "rgba(16,185,129,0.4)"),
    "Medium": ("#f97316", "rgba(249,115,22,0.15)",  "rgba(249,115,22,0.4)"),
    "Hard":   ("#ef4444", "rgba(239,68,68,0.15)",   "rgba(239,68,68,0.4)"),
}
_DIFF_ICONS = {"Easy": "🟢", "Medium": "🟠", "Hard": "🔴"}

# minutes per question per difficulty
_DIFF_MINS  = {"Easy": 2, "Medium": 4, "Hard": 6}


def _question_card(idx: int, question: str, accent: str,
                   badge_label: str = "", badge_color: str = "") -> None:
    badge_html = ""
    if badge_label:
        b_clr, b_bg, b_border = _DIFF_COLORS.get(
            badge_label, ("#94a3b8", "rgba(148,163,184,0.15)", "rgba(148,163,184,0.4)")
        )
        icon = _DIFF_ICONS.get(badge_label, "")
        badge_html = (
            f"<span style='background:{b_bg};color:{b_clr};border:1px solid {b_border};"
            f"padding:2px 10px;border-radius:20px;font-size:0.65rem;font-weight:700;"
            f"margin-left:8px'>{icon} {badge_label}</span>"
        )

    st.markdown(
        f"""
        <div style="background:rgba(255,255,255,0.04);border-radius:14px;
                    padding:16px 20px;border:1px solid {accent}25;
                    border-left:3px solid {accent};margin-bottom:10px">
            <div style="display:flex;align-items:center;margin-bottom:10px">
                <span style="background:{accent}20;color:{accent};font-weight:800;
                             font-size:0.72rem;padding:3px 10px;border-radius:20px">
                    Q{idx}</span>
                {badge_html}
            </div>
            <div style="font-size:0.88rem;color:#f1f5f9;font-weight:500;line-height:1.7">
                {html.escape(question)}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


# ── PDF / DOCX builders ───────────────────────────────────────────────────

def _build_pdf(job_title: str, cand_name: str,
               tech_qs: list, beh_qs: list, sit_qs: list) -> bytes:
    from fpdf import FPDF
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Interview Questions", ln=True, align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 7, f"Job: {job_title}   |   Candidate: {cand_name}", ln=True, align="C")
    pdf.ln(6)
    for sec_title, items in [("Technical Questions", tech_qs),
                              ("Behavioral Questions", beh_qs),
                              ("Situational Questions", sit_qs)]:
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_fill_color(230, 230, 250)
        pdf.cell(0, 9, sec_title, ln=True, fill=True)
        pdf.ln(2)
        for i, item in enumerate(items, 1):
            diff = item.get("difficulty", "")
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 7, f"Q{i}  [{diff}]" if diff else f"Q{i}", ln=True)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 6, item.get("question", ""))
            pdf.ln(2)
        pdf.ln(3)
    return bytes(pdf.output())


def _build_docx(job_title: str, cand_name: str,
                tech_qs: list, beh_qs: list, sit_qs: list) -> bytes:
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    t = doc.add_heading("Interview Questions", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Job: {job_title}   |   Candidate: {cand_name}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    level_colors = {
        "Easy":   RGBColor(0x10, 0xb9, 0x81),
        "Medium": RGBColor(0xf9, 0x73, 0x16),
        "Hard":   RGBColor(0xef, 0x44, 0x44),
    }
    for sec_title, items in [("Technical Questions", tech_qs),
                              ("Behavioral Questions", beh_qs),
                              ("Situational Questions", sit_qs)]:
        doc.add_heading(sec_title, level=1)
        for i, item in enumerate(items, 1):
            diff = item.get("difficulty", "")
            p = doc.add_paragraph()
            r = p.add_run(f"Q{i}  ")
            r.bold = True
            r.font.size = Pt(11)
            if diff:
                rb = p.add_run(f"[{diff}]  ")
                rb.bold = True
                rb.font.size = Pt(9)
                rb.font.color.rgb = level_colors.get(diff, RGBColor(0x64, 0x74, 0x8b))
            rq = p.add_run(item.get("question", ""))
            rq.font.size = Pt(11)
        doc.add_paragraph()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Difficulty-grouped renderer ──────────────────────────────────────────────

_DIFF_ORDER = ["Easy", "Medium", "Hard"]

_DIFF_SECTION = {
    "Easy":   ("🟢 Easy",   "#10b981", "rgba(16,185,129,0.08)",  "rgba(16,185,129,0.25)"),
    "Medium": ("🟠 Medium", "#f97316", "rgba(249,115,22,0.08)",  "rgba(249,115,22,0.25)"),
    "Hard":   ("🔴 Hard",   "#ef4444", "rgba(239,68,68,0.08)",   "rgba(239,68,68,0.25)"),
}


def _est_time(items: list[dict]) -> int:
    return sum(_DIFF_MINS.get(i.get("difficulty", "Medium"), 4) for i in items)


def _render_by_difficulty(items: list[dict], accent: str, empty_msg: str) -> None:
    """Group questions by difficulty and render each group with a colored header + time estimate."""
    if not items:
        st.info(empty_msg)
        return

    for item in items:
        if item.get("difficulty") not in _DIFF_ORDER:
            item["difficulty"] = "Medium"

    grouped: dict[str, list[dict]] = {d: [] for d in _DIFF_ORDER}
    for item in items:
        grouped[item["difficulty"]].append(item)

    total_mins = _est_time(items)
    st.markdown(
        f"<div style='font-size:0.72rem;color:#94a3b8;margin-bottom:14px'>"
        f"⏱️ Estimated section time: <b style='color:#e2e8f0'>{total_mins} min</b></div>",
        unsafe_allow_html=True,
    )

    global_idx = 1
    for diff in _DIFF_ORDER:
        bucket = grouped[diff]
        if not bucket:
            continue
        label, clr, bg, border = _DIFF_SECTION[diff]
        mins = _est_time(bucket)
        st.markdown(
            f"<div style='background:{bg};border:1px solid {border};"
            f"border-radius:10px;padding:8px 16px;margin:18px 0 10px 0;"
            f"display:flex;align-items:center;gap:16px'>"
            f"<span style='color:{clr};font-size:0.78rem;font-weight:800;"
            f"letter-spacing:0.06em'>{label}</span>"
            f"<span style='color:{clr};font-size:0.72rem;font-weight:600;opacity:0.8'>"
            f"{len(bucket)} question{'s' if len(bucket)>1 else ''}</span>"
            f"<span style='color:{clr};font-size:0.7rem;opacity:0.7;margin-left:auto'>"
            f"⏱️ ~{mins} min</span>"
            f"</div>",
            unsafe_allow_html=True,
        )
        for item in bucket:
            _question_card(global_idx, item.get("question", ""), accent, badge_label=diff)
            global_idx += 1


# ── Main render ────────────────────────────────────────────────────────────

def render(service: CandidateService) -> None:
    page_header(
        "❓ Interview Questions",
        "Generate role-specific interview questions using Groq AI based on JD and candidate profile.",
    )

    candidates = service.get_all_candidates(st.session_state.get("recruiter_email", ""))
    jobs       = service.jobs.get_all_jobs(st.session_state.get("recruiter_email", ""))

    if not candidates:
        empty_state("No candidates found — upload resumes first.")
        return
    if not jobs:
        empty_state("No jobs found — create a job posting first.")
        return

    # ── Selectors ──────────────────────────────────────────────────────────
    with st.container(border=True):
        st.markdown(
            "<div style='font-size:0.8rem;font-weight:700;color:#94a3b8;"
            "text-transform:uppercase;letter-spacing:0.08em;margin-bottom:14px'>"
            "⚙️ Configuration</div>",
            unsafe_allow_html=True,
        )
        col1, col2 = st.columns(2)
        with col1:
            job_options = {f"#{j['job_id']} — {j['job_title']}": j for j in jobs}
            sel_job_key = st.selectbox("💼 Job Description", list(job_options.keys()), key="iq_job")
            sel_job     = job_options[sel_job_key]
        with col2:
            cand_options = {
                f"#{c['candidate_id']} — {(c.get('name') or 'Unknown').splitlines()[0]}": c
                for c in candidates
            }
            sel_cand_key = st.selectbox("👤 Candidate", list(cand_options.keys()), key="iq_cand")
            sel_cand     = cand_options[sel_cand_key]

        generate = st.button("🚀 Generate Questions", type="primary", use_container_width=True)

    if not generate and "iq_extracted" not in st.session_state:
        st.info("👆 Select a Job Description and Candidate, then click **Generate Questions**.")
        return

    # ── Extract fields ─────────────────────────────────────────────────────
    if generate:
        jd_skills           = [s.strip() for s in (sel_job.get("skills_required") or "").split(",") if s.strip()]
        jd_responsibilities = [r.strip() for r in (sel_job.get("responsibilities") or "").splitlines() if r.strip()]
        cand_skills         = [s.strip() for s in (sel_cand.get("skills") or "").split(",") if s.strip()]
        cand_exp            = [e.strip() for e in (sel_cand.get("experience") or "").splitlines() if e.strip()]
        jd_set              = {s.lower() for s in jd_skills}
        c_set               = {s.lower() for s in cand_skills}

        st.session_state.iq_extracted = {
            "job":               sel_job,
            "candidate":         sel_cand,
            "jd_skills":         jd_skills,
            "jd_responsibilities": jd_responsibilities,
            "cand_skills":       cand_skills,
            "cand_exp":          cand_exp,
            "matched":           sorted(jd_set & c_set),
            "missing":           sorted(jd_set - c_set),
        }
        st.session_state.pop("iq_questions", None)

    data      = st.session_state.iq_extracted
    job       = data["job"]
    cand      = data["candidate"]
    cand_name = (cand.get("name") or "Unknown").splitlines()[0].strip()
    job_title = job.get("job_title") or "—"

    st.divider()

    # ── Extracted info panel ───────────────────────────────────────────────
    st.markdown(
        "<div style='font-size:0.9rem;font-weight:700;color:#e2e8f0;margin-bottom:16px'>"
        "📋 Extracted Information</div>",
        unsafe_allow_html=True,
    )

    col_jd, col_cand = st.columns(2)

    with col_jd:
        with st.container(border=True):
            st.markdown(
                "<div style='font-size:0.72rem;font-weight:700;color:#60a5fa;"
                "text-transform:uppercase;letter-spacing:0.1em;margin-bottom:14px'>"
                "💼 Job Description</div>",
                unsafe_allow_html=True,
            )
            _info_row("Job Title",   job_title)
            _info_row("Department",  job.get("department") or "—")
            _info_row("Experience",  job.get("experience_level") or "—")
            _info_row("Job Type",    job.get("job_type") or "—")
            st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
            _section_label("🛠 Required Skills", "#a78bfa")
            if data["jd_skills"]:
                st.markdown(
                    f"<div style='line-height:2.4'>{_pills(data['jd_skills'], '#c4b5fd', 'rgba(139,92,246,0.15)', 'rgba(139,92,246,0.4)')}</div>",
                    unsafe_allow_html=True,
                )
            else:
                st.caption("No skills listed.")
            st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
            _section_label("📌 Responsibilities", "#60a5fa")
            if data["jd_responsibilities"]:
                for r in data["jd_responsibilities"][:5]:
                    st.markdown(
                        f"<div style='font-size:0.8rem;color:#cbd5e1;padding:4px 0 4px 10px;"
                        f"border-left:2px solid rgba(96,165,250,0.4);margin-bottom:4px'>"
                        f"{html.escape(r)}</div>",
                        unsafe_allow_html=True,
                    )
            else:
                st.caption("No responsibilities listed.")

    with col_cand:
        with st.container(border=True):
            st.markdown(
                "<div style='font-size:0.72rem;font-weight:700;color:#34d399;"
                "text-transform:uppercase;letter-spacing:0.1em;margin-bottom:14px'>"
                "👤 Candidate Profile</div>",
                unsafe_allow_html=True,
            )
            _info_row("Name",  cand_name)
            _info_row("Email", cand.get("email") or "—")
            _info_row("Phone", cand.get("phone") or "—")
            st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
            _section_label("🛠 Candidate Skills", "#34d399")
            if data["cand_skills"]:
                st.markdown(
                    f"<div style='line-height:2.4'>{_pills(data['cand_skills'], '#6ee7b7', 'rgba(16,185,129,0.15)', 'rgba(16,185,129,0.4)')}</div>",
                    unsafe_allow_html=True,
                )
            else:
                st.caption("No skills found.")
            st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
            _section_label("💼 Experience", "#fbbf24")
            if data["cand_exp"]:
                for e in data["cand_exp"][:5]:
                    st.markdown(
                        f"<div style='font-size:0.8rem;color:#cbd5e1;padding:4px 0 4px 10px;"
                        f"border-left:2px solid rgba(251,191,36,0.4);margin-bottom:4px'>"
                        f"{html.escape(e)}</div>",
                        unsafe_allow_html=True,
                    )
            else:
                st.caption("No experience found.")

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Skill comparison ───────────────────────────────────────────────────
    with st.container(border=True):
        st.markdown(
            "<div style='font-size:0.72rem;font-weight:700;color:#94a3b8;"
            "text-transform:uppercase;letter-spacing:0.1em;margin-bottom:14px'>"
            "🔍 Skill Comparison</div>",
            unsafe_allow_html=True,
        )
        total_jd  = len(data["jd_skills"])
        matched_n = len(data["matched"])
        missing_n = len(data["missing"])
        match_pct = round(matched_n / total_jd * 100) if total_jd else 0

        m1, m2, m3, m4 = st.columns(4)
        m1.metric("📋 JD Skills",  total_jd)
        m2.metric("✅ Matched",    matched_n)
        m3.metric("❌ Missing",    missing_n)
        m4.metric("🎯 Match %",    f"{match_pct}%")

        st.markdown(
            f"<div style='background:rgba(255,255,255,0.08);border-radius:20px;"
            f"height:8px;overflow:hidden;margin:12px 0'>"
            f"<div style='background:linear-gradient(90deg,#10b981,#34d399);"
            f"height:100%;width:{match_pct}%;border-radius:20px'></div></div>",
            unsafe_allow_html=True,
        )
        sc1, sc2 = st.columns(2)
        with sc1:
            _section_label("✅ Matched Skills", "#10b981")
            st.markdown(
                f"<div style='line-height:2.4'>{_pills(data['matched'], '#6ee7b7', 'rgba(16,185,129,0.15)', 'rgba(16,185,129,0.4)') or '<span style=\"color:#64748b;font-size:0.8rem\">None</span>'}</div>",
                unsafe_allow_html=True,
            )
        with sc2:
            _section_label("❌ Skill Gaps", "#ef4444")
            st.markdown(
                f"<div style='line-height:2.4'>{_pills(data['missing'], '#fca5a5', 'rgba(239,68,68,0.15)', 'rgba(239,68,68,0.4)') or '<span style=\"color:#64748b;font-size:0.8rem\">None — perfect match!</span>'}</div>",
                unsafe_allow_html=True,
            )

    st.divider()

    # ── Groq generation ────────────────────────────────────────────────────
    if "iq_questions" not in st.session_state:
        if not GROQ_API_KEY:
            st.error("⚠️ Groq API key not set. Add your key to `config/settings.py` → `GROQ_API_KEY`.")
            return

        with st.spinner("🤖 Generating questions with Groq AI..."):
            try:
                qs = _generate_via_groq(
                    job_title        = job_title,
                    jd_skills        = data["jd_skills"],
                    responsibilities = data["jd_responsibilities"],
                    cand_name        = cand_name,
                    cand_skills      = data["cand_skills"],
                    cand_exp         = data["cand_exp"],
                    matched          = data["matched"],
                    missing          = data["missing"],
                )
                st.session_state.iq_questions = qs
            except RuntimeError as e:
                st.warning(f"⏳ {e}")
                return
            except Exception as e:
                st.error(f"❌ Groq generation failed: {e}")
                return

    qs = st.session_state.iq_questions

    # ── Question tabs ──────────────────────────────────────────────────────
    st.markdown(
        "<div style='font-size:0.9rem;font-weight:700;color:#e2e8f0;margin-bottom:16px'>"
        "📝 Generated Interview Questions</div>",
        unsafe_allow_html=True,
    )

    tech_qs = qs.get("technical", [])
    beh_qs  = qs.get("behavioral", [])
    sit_qs  = qs.get("situational", [])

    tab1, tab2, tab3 = st.tabs([
        f"🛠 Technical ({len(tech_qs)}) ⏱️ ~{_est_time(tech_qs)}min",
        f"🧠 Behavioral ({len(beh_qs)}) ⏱️ ~{_est_time(beh_qs)}min",
        f"💡 Situational ({len(sit_qs)}) ⏱️ ~{_est_time(sit_qs)}min",
    ])

    with tab1:
        _render_by_difficulty(tech_qs, "#8b5cf6", "No technical questions generated.")

    with tab2:
        _render_by_difficulty(beh_qs, "#3b82f6", "No behavioral questions generated.")

    with tab3:
        _render_by_difficulty(sit_qs, "#10b981", "No situational questions generated.")

    st.divider()

    # ── Download buttons ───────────────────────────────────────────────────
    dl1, dl2 = st.columns(2)
    with dl1:
        st.download_button(
            label="📄 Download as PDF",
            data=_build_pdf(job_title, cand_name, tech_qs, beh_qs, sit_qs),
            file_name=f"interview_questions_{cand_name.replace(' ', '_')}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    with dl2:
        st.download_button(
            label="📝 Download as DOCX",
            data=_build_docx(job_title, cand_name, tech_qs, beh_qs, sit_qs),
            file_name=f"interview_questions_{cand_name.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
